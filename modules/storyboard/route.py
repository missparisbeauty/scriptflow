"""分鏡 route — Phase 5。

POST /api/v1/storyboard/generate
GET  /api/v1/storyboard/{id}/export?format=pdf|word

注意：generate 回傳的 scenes[].image_bytes 是 raw bytes，
直接 JSON 序列化會失敗，這層需要轉 base64 或剝掉。
"""

from __future__ import annotations

import base64
from io import BytesIO

from fastapi import APIRouter, Depends, Path, Query
from fastapi.responses import StreamingResponse

from domain.exceptions import InvalidInput
from domain.responses import ok
from modules.auth.middleware import require_session
from modules.storyboard.schema import (
    ExportFormatLiteral,
    GenerateRequest,
    PlatformLiteral,
)
from modules.storyboard.service import generate as sb_generate

router = APIRouter(
    prefix="/api/v1/storyboard",
    tags=["storyboard"],
    dependencies=[Depends(require_session)],
)


@router.post("/generate")
def generate_storyboard(req: GenerateRequest) -> dict:
    raw = sb_generate(req.script_id, req.platform)
    # bytes 轉 base64 才能 JSON 序列化
    raw["scenes"] = [_serialize_scene(s) for s in raw["scenes"]]
    return ok(raw)


@router.get("/{storyboard_id}/export")
def export_storyboard(
    storyboard_id: str = Path(..., min_length=3, max_length=80),
    format: ExportFormatLiteral = Query("word"),
    # storyboard 不持久化，這個端點需要前端帶 script_id + platform
    script_id: str = Query(..., min_length=1, max_length=80),
    platform: PlatformLiteral = Query("ig_reels"),
) -> StreamingResponse:
    """重新產一次 storyboard 並匯出。

    註：因為 Phase 4 設計 storyboard 不存 collection，匯出端點
    需要 script_id + platform 重算。Phase 7 可改持久化。
    """
    raw = sb_generate(script_id, platform)

    if format == "word":
        buf = _export_word(raw)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:  # pdf - 暫不實作真正的 PDF（Phase 7）
        raise InvalidInput("pdf export not implemented yet; use word")

    return StreamingResponse(
        BytesIO(buf),
        media_type=media,
        headers={
            "Content-Disposition": f'attachment; filename="storyboard_{storyboard_id}.{ext}"',
        },
    )


# --- 內部 helper ---


def _serialize_scene(scene: dict) -> dict:
    """把 image_bytes 轉成 base64 data URL，方便前端 <img src=...>。"""
    out = dict(scene)
    img = out.pop("image_bytes", None)
    if img:
        out["image_data_url"] = (
            "data:image/png;base64," + base64.b64encode(img).decode("ascii")
        )
    else:
        out["image_data_url"] = None
    return out


def _export_word(raw: dict) -> bytes:
    """匯出 .docx；用 python-docx，不放 secret 或外部資料。"""
    from docx import Document

    doc = Document()
    doc.add_heading(f"Storyboard {raw.get('storyboard_id', '')}", level=1)
    doc.add_paragraph(
        f"script_id: {raw.get('script_id', '')} | platform: {raw.get('platform', '')}"
    )
    for scene in raw.get("scenes", []):
        doc.add_heading(
            f"鏡頭 {scene.get('index', '?')}（{scene.get('time', '')}）", level=2
        )
        doc.add_paragraph(f"畫面：{scene.get('scene', '')}")
        doc.add_paragraph(f"口白：{scene.get('voiceover', '')}")
        doc.add_paragraph(f"音效：{scene.get('sfx', '')}")
        doc.add_paragraph(f"產品露出策略：{scene.get('product_exposure', '')}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
